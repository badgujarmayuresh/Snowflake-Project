"""
Generates project_proposal.docx for Smart BI Agent project.
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from pathlib import Path

OUT = Path(__file__).parent.parent / "project_proposal.docx"

# ── Colour helpers ──────────────────────────────────────────────
SNOW_BLUE = "29B5E8"
DARK_BG   = "0D1B2A"
ACCENT    = "00C4B4"
WHITE     = "FFFFFF"
GREY_BG   = "F2F4F8"
TEXT_DARK = "1A1A2E"
RED       = "E83A3A"

def set_cell_bg(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)

def set_cell_border(cell, hex_color="CCCCCC", size="4"):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ("top","left","bottom","right"):
        border = OxmlElement(f"w:{side}")
        border.set(qn("w:val"),   "single")
        border.set(qn("w:sz"),    size)
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), hex_color)
        tcBorders.append(border)
    tcPr.append(tcBorders)

def add_para(doc, text, size=11, bold=False, color=TEXT_DARK,
             align=WD_ALIGN_PARAGRAPH.LEFT, space_before=0, space_after=6, italic=False):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    run = p.add_run(text)
    run.font.size  = Pt(size)
    run.font.bold  = bold
    run.font.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)
    return p

def add_heading(doc, text, level=1):
    sizes   = {1: 22, 2: 16, 3: 13}
    colors  = {1: DARK_BG, 2: SNOW_BLUE, 3: DARK_BG}
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run(text)
    run.font.size  = Pt(sizes.get(level, 12))
    run.font.bold  = True
    run.font.color.rgb = RGBColor.from_string(colors.get(level, TEXT_DARK))
    if level == 1:
        # bottom border
        pPr  = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bot  = OxmlElement("w:bottom")
        bot.set(qn("w:val"),   "single")
        bot.set(qn("w:sz"),    "6")
        bot.set(qn("w:space"), "1")
        bot.set(qn("w:color"), SNOW_BLUE)
        pBdr.append(bot)
        pPr.append(pBdr)
    return p

def add_table(doc, headers, rows, col_widths=None, header_bg=DARK_BG):
    table = doc.add_table(rows=1+len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"

    # header row
    hdr_row = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr_row.cells[i]
        cell.text = h
        set_cell_bg(cell, header_bg)
        for run in cell.paragraphs[0].runs:
            run.font.bold  = True
            run.font.color.rgb = RGBColor.from_string(WHITE)
            run.font.size  = Pt(10)

    # data rows
    for r_idx, row in enumerate(rows):
        bg = WHITE if r_idx % 2 == 0 else GREY_BG
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx+1].cells[c_idx]
            cell.text = str(val)
            set_cell_bg(cell, bg)
            for run in cell.paragraphs[0].runs:
                run.font.size = Pt(10)
            set_cell_border(cell)

    # column widths
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Inches(w)
    return table

def add_bullet(doc, text, level=0, color=TEXT_DARK):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after  = Pt(3)
    p.paragraph_format.left_indent  = Inches(0.25 * (level+1))
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor.from_string(color)

# ══════════════════════════════════════════════════════════════
doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin    = Cm(1.8)
    section.bottom_margin = Cm(1.8)
    section.left_margin   = Cm(2.2)
    section.right_margin  = Cm(2.2)

# ── COVER ──────────────────────────────────────────────────────
add_para(doc, "", space_before=30)

title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title_p.add_run("Smart BI Agent")
r.font.size  = Pt(36)
r.font.bold  = True
r.font.color.rgb = RGBColor.from_string(DARK_BG)

sub_p = doc.add_paragraph()
sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub_p.add_run("Intelligent Analytics Platform for Retail & Manufacturing")
r.font.size  = Pt(16)
r.font.color.rgb = RGBColor.from_string(SNOW_BLUE)

add_para(doc, "", space_before=8)

meta_p = doc.add_paragraph()
meta_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = meta_p.add_run("Prepared by: Mayuresh Prakash Badgujar  |  May 2026  |  Version 1.0")
r.font.size  = Pt(11)
r.font.color.rgb = RGBColor.from_string("555555")

add_para(doc, "", space_before=4)

tag_p = doc.add_paragraph()
tag_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = tag_p.add_run("Powered by Snowflake Cortex AI")
r.font.size   = Pt(11)
r.font.italic = True
r.font.color.rgb = RGBColor.from_string(ACCENT)

doc.add_page_break()

# ── 1. EXECUTIVE SUMMARY ───────────────────────────────────────
add_heading(doc, "1. Executive Summary")
add_para(doc,
    "Organizations today sit on mountains of data but struggle to extract timely, actionable insights. "
    "Analysts are bottlenecked. Business users wait days for answers. Decisions are made on instinct rather than evidence.",
    space_after=8)
add_para(doc,
    "Smart BI Agent solves this by combining 10+ years of BI & Analytics expertise with Snowflake's "
    "cutting-edge AI capabilities — delivering a platform where any business user can ask questions "
    "in plain English and get instant, accurate, data-driven answers.",
    space_after=8)

# ── 2. THE PROBLEM ─────────────────────────────────────────────
add_heading(doc, "2. The Problem")
add_table(doc,
    ["Pain Point", "Business Impact"],
    [
        ["No Self-Service",    "Only analysts can query data — business waits days for answers"],
        ["Buried Insights",    "Reports exist but no one can find or search through them"],
        ["No Predictions",     "Teams react to problems instead of preventing them"],
        ["Fragmented Tools",   "BI + ML + Apps are separate — high cost, inconsistent data"],
    ],
    col_widths=[2.5, 4.5])

# ── 3. THE SOLUTION ────────────────────────────────────────────
add_heading(doc, "3. The Solution")
add_para(doc,
    "A single, unified intelligent analytics platform built entirely on Snowflake that combines:",
    space_after=4)
add_bullet(doc, "Natural Language Analytics — Ask questions, get answers. No SQL required.")
add_bullet(doc, "Semantic Search — Find insights from historical reports instantly.")
add_bullet(doc, "Predictive Intelligence — Forecast demand and detect anomalies automatically.")
add_bullet(doc, "One App — Everything delivered through a single, simple Streamlit interface.")

# ── 4. DATA DOMAIN ─────────────────────────────────────────────
add_heading(doc, "4. Data Domain")
add_para(doc,
    "The platform covers a company that manufactures products and sells them online — end to end.",
    space_after=6)
add_table(doc,
    ["Domain", "Coverage", "Key Entities"],
    [
        ["E-commerce / Retail",  "Online retail and customer side",    "Orders, Customers, Payments, Reviews, Products"],
        ["B2B / Wholesale",      "Supplier and employee sales",        "Suppliers, Product Catalogue, Wholesale Orders"],
        ["Manufacturing",        "Factory and production operations",  "Machines, Production Orders, Defects, Inventory"],
    ],
    col_widths=[2.0, 2.5, 3.0])

# ── 5. ARCHITECTURE ────────────────────────────────────────────
add_heading(doc, "5. Layered Architecture")
add_para(doc,
    "The platform follows a Bronze / Silver / Gold layered architecture — a modern data engineering "
    "best practice that ensures raw data is always preserved, transformations are incremental and auditable, "
    "and the Gold layer is always business-ready.",
    space_after=8)

add_heading(doc, "5.1  Bronze Layer — Raw Ingestion", level=2)
add_bullet(doc, "Data loaded as-is from CSV files using Snowflake internal stages (COPY INTO)")
add_bullet(doc, "No transformations or business logic applied")
add_bullet(doc, "Acts as the system of record — full replayability if anything breaks downstream")
add_bullet(doc, "Schema: DB_DEMO_MAYURESH.BRONZE")

add_heading(doc, "5.2  Silver Layer — Cleansing & Standardisation", level=2)
add_bullet(doc, "Null handling — replace or flag nulls with meaningful defaults")
add_bullet(doc, "Type casting — ensure dates, amounts, and IDs are correctly typed")
add_bullet(doc, "Column renaming — standard snake_case naming convention")
add_bullet(doc, "Deduplication — remove duplicates using ROW_NUMBER()")
add_bullet(doc, "Schema: DB_DEMO_MAYURESH.SILVER")

add_heading(doc, "5.3  Gold Layer — Dimensional Model", level=2)
add_bullet(doc, "Business-ready star schema optimised for analytics and AI consumption")
add_bullet(doc, "Fact tables: FACT_SALES, FACT_ORDERS, FACT_PRODUCTION, FACT_DEFECTS")
add_bullet(doc, "Dimension tables: DIM_DATE, DIM_CUSTOMER, DIM_PRODUCT, DIM_SELLER, DIM_MACHINE, DIM_LOCATION")
add_bullet(doc, "Schema: DB_DEMO_MAYURESH.GOLD")

# ── 6. DIMENSIONAL MODEL ───────────────────────────────────────
add_heading(doc, "6. Dimensional Model — Star Schema")

add_heading(doc, "Fact Tables", level=3)
add_table(doc,
    ["Fact Table", "Grain", "Key Measures"],
    [
        ["FACT_SALES",       "One row per order item",       "Revenue, quantity, freight, discount"],
        ["FACT_ORDERS",      "One row per order",            "Order count, payment value, delivery days"],
        ["FACT_PRODUCTION",  "One row per production order", "Planned vs actual quantity, production cost"],
        ["FACT_DEFECTS",     "One row per defect record",    "Defect quantity, defect rate"],
    ],
    col_widths=[2.0, 2.3, 3.2])

add_para(doc, "", space_after=6)
add_heading(doc, "Dimension Tables", level=3)
add_table(doc,
    ["Dimension Table", "Description"],
    [
        ["DIM_DATE",      "Calendar — year, quarter, month, week, day"],
        ["DIM_CUSTOMER",  "Customer details, city, state, region"],
        ["DIM_PRODUCT",   "Product, category, translated category name"],
        ["DIM_SELLER",    "Seller details and location"],
        ["DIM_MACHINE",   "Machine type, plant location, status, capacity"],
        ["DIM_LOCATION",  "Full geographic hierarchy"],
    ],
    col_widths=[2.5, 5.0])

# ── 7. AI CAPABILITIES ─────────────────────────────────────────
add_heading(doc, "7. AI Capabilities")

add_heading(doc, "7.1  Cortex Analyst — Natural Language to SQL", level=2)
add_para(doc,
    "Cortex Analyst uses a semantic YAML model built on top of the Gold layer to translate "
    "business questions into SQL — no coding required.",
    space_after=4)
add_bullet(doc, 'Example: "What were the top 10 products by revenue last quarter?"')
add_bullet(doc, 'Example: "Show me monthly sales trend for electronics."')
add_bullet(doc, 'Example: "Which plant had the highest defect rate in 2024?"')

add_heading(doc, "7.2  Cortex Search — Semantic Insight Discovery", level=2)
add_para(doc,
    "Cortex Search indexes key business insights, report summaries, and anomaly descriptions — "
    "enabling semantic search across historical context.",
    space_after=4)
add_bullet(doc, 'Example: "Find all reports about sales decline."')
add_bullet(doc, 'Example: "What happened with machine downtime in Plant A last year?"')

add_heading(doc, "7.3  Cortex ML — Forecasting & Anomaly Detection", level=2)
add_para(doc,
    "Built-in ML functions that require no data science expertise.",
    space_after=4)
add_bullet(doc, "Forecasting — predict future sales and production demand")
add_bullet(doc, "Anomaly Detection — automatically flag unusual patterns in sales, defects, or downtime")
add_bullet(doc, "Classification — classify customer churn risk and product quality")

add_heading(doc, "7.4  Snowflake Agent — AI Orchestration", level=2)
add_para(doc,
    "The Snowflake Agent acts as an intelligent orchestrator — deciding which AI tool to use "
    "based on the user's question and combining results into a single, coherent response.",
    space_after=4)

# ── 8. TECH STACK ──────────────────────────────────────────────
add_heading(doc, "8. Technology Stack")
add_table(doc,
    ["Layer", "Technology", "Purpose"],
    [
        ["Cloud Platform",    "Snowflake",               "Single platform — storage, compute, AI, app"],
        ["Data Pipeline",     "Bronze / Silver / Gold",   "Structured, layered data architecture"],
        ["Dimensional Model", "Star Schema",              "Optimised for analytics and AI"],
        ["Natural Language",  "Cortex Analyst",           "Business questions translated to SQL"],
        ["Semantic Search",   "Cortex Search Service",    "Index and retrieve insights semantically"],
        ["Predictive AI",     "Cortex ML Functions",      "Built-in forecasting and anomaly detection"],
        ["Orchestration",     "Snowflake Cortex Agent",   "Intelligent routing across all AI tools"],
        ["User Interface",    "Streamlit in Snowflake",   "Native app — no infrastructure required"],
    ],
    col_widths=[1.8, 2.2, 3.5])

# ── 9. BUSINESS VALUE ──────────────────────────────────────────
add_heading(doc, "9. Business Value")
add_table(doc,
    ["Metric", "Before", "After"],
    [
        ["Time to answer a question",  "Hours / Days",   "Seconds"],
        ["Users who can self-serve",   "Analysts only",  "Everyone"],
        ["Forecast availability",      "Manual / Excel", "Automated, always current"],
        ["Anomaly detection",          "Reactive",       "Proactive, real-time"],
        ["Number of tools required",   "Multiple",       "One (Snowflake)"],
    ],
    col_widths=[2.8, 2.0, 2.7])

# ── 10. DELIVERY PHASES ────────────────────────────────────────
add_heading(doc, "10. Delivery Phases")
add_table(doc,
    ["Phase", "Deliverable", "Description"],
    [
        ["1", "Architecture & Design",  "Define layers, data model, and tech decisions"],
        ["2", "Bronze Layer",           "Ingest raw CSV data into Snowflake as-is"],
        ["3", "Silver Layer",           "Clean, standardise, and type-cast data"],
        ["4", "Gold Layer",             "Build dimensional model — facts + dimensions"],
        ["5", "Cortex Analyst",         "Natural language interface on Gold layer"],
        ["6", "Cortex Search",          "Semantic search across historical insights"],
        ["7", "Cortex ML",              "Forecasting and anomaly detection"],
        ["8", "Snowflake Agent",        "Orchestrate all AI tools intelligently"],
        ["9", "Streamlit App",          "Business user interface — all in one place"],
    ],
    col_widths=[0.7, 2.0, 4.8])

# ── 11. ABOUT ──────────────────────────────────────────────────
add_heading(doc, "11. About This Project")
add_para(doc,
    "This platform is built by a BI & Analytics professional with 10+ years of experience across "
    "ETL, data modelling, and reporting — now extended with modern AI capabilities on Snowflake.",
    space_after=6)
add_para(doc,
    "It serves as both a working production platform and a reference architecture for how traditional "
    "BI teams can evolve into AI-powered analytics organisations.",
    space_after=6)
add_para(doc,
    "Built with Snowflake Cortex AI  |  Streamlit  |  Star Schema Dimensional Modelling",
    size=10, italic=True, color="888888", align=WD_ALIGN_PARAGRAPH.CENTER, space_before=20)

doc.save(str(OUT))
print(f"Saved: {OUT}")
