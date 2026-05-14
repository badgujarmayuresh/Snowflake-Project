from docx import Document

doc = Document("Naming Conventions and Snowflake Standards.docx")

for para in doc.paragraphs:
    if para.text.strip():
        print(para.text)

# Also read tables if any
for table in doc.tables:
    for row in table.rows:
        row_data = [cell.text.strip() for cell in row.cells]
        if any(row_data):
            print(" | ".join(row_data))
